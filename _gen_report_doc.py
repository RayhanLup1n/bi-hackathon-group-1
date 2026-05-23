"""
Generate progress-report.docx from structured data.
Run: python _gen_report_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color Palette ─────────────────────────────────────────
PRIMARY   = RGBColor(0x1a, 0x3a, 0x5c)
ACCENT    = RGBColor(0xe8, 0x5d, 0x04)
GREEN     = RGBColor(0x2d, 0x6a, 0x4f)
YELLOW    = RGBColor(0xb5, 0x83, 0x0a)
RED_COL   = RGBColor(0xc0, 0x39, 0x2b)
MUTED     = RGBColor(0x6c, 0x75, 0x7d)
WHITE     = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG  = RGBColor(0xf8, 0xf9, 0xfa)
BG_GREEN  = RGBColor(0xd8, 0xf3, 0xdc)
BG_YELLOW = RGBColor(0xff, 0xf3, 0xcd)
BG_RED    = RGBColor(0xfd, 0xe8, 0xe8)
BG_BLUE   = RGBColor(0xe8, 0xf0, 0xfb)

doc = Document()

# ── Page margins ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'),   kwargs[side].get('val', 'single'))
            tag.set(qn('w:sz'),    kwargs[side].get('sz', '4'))
            tag.set(qn('w:space'), kwargs[side].get('space', '0'))
            tag.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def paragraph_space(doc, before=0, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    # Bottom border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a3a5c')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = PRIMARY
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = PRIMARY
    return p

def body_para(doc, text, indent=False, italic=False, color=None, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(9.5)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
    return p

def alert_box(doc, text, color: RGBColor, border_color: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    set_cell_border(cell,
        top    ={'val':'single','sz':'12','color': border_color},
        left   ={'val':'single','sz':'24','color': border_color},
        bottom ={'val':'single','sz':'12','color': border_color},
        right  ={'val':'single','sz':'12','color': border_color},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

STATUS = {
    '✅ Selesai':             (BG_GREEN,  GREEN),
    '🔄 Perlu Update':        (BG_YELLOW, YELLOW),
    '🔄 Sebagian':            (BG_YELLOW, YELLOW),
    '🔄 Perlu Angka':         (BG_YELLOW, YELLOW),
    '🔄 Perlu Diperketat':    (BG_YELLOW, YELLOW),
    '❌ Belum':               (BG_RED,    RED_COL),
    '❌ Belum terkonfirmasi': (BG_RED,    RED_COL),
    '❓ Ada, belum diverifikasi': (LIGHT_BG, MUTED),
    '🔄 Ada, sebagian':       (BG_YELLOW, YELLOW),
    '⚠ WAJIB per proposal':  (BG_RED,    RED_COL),
}

def make_status_table(doc, rows_data, col_widths=None):
    """
    rows_data = list of (komponen, status_str, catatan)
    """
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hdr = tbl.rows[0].cells
    labels = ['Komponen', 'Status', 'Catatan']
    hdr_widths = col_widths or [Cm(8), Cm(3.5), Cm(5.5)]
    for i, (cell, label, w) in enumerate(zip(hdr, labels, hdr_widths)):
        set_cell_bg(cell, PRIMARY)
        cell.width = w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    for idx, (comp, status, note) in enumerate(rows_data):
        row = tbl.add_row()
        bg = LIGHT_BG if idx % 2 == 1 else RGBColor(0xff,0xff,0xff)
        for c in row.cells:
            set_cell_bg(c, bg)

        # Col 0 — component
        p = row.cells[0].paragraphs[0]
        run = p.add_run(comp)
        run.font.size = Pt(9)

        # Col 1 — status badge
        s_bg, s_fg = STATUS.get(status, (LIGHT_BG, MUTED))
        set_cell_bg(row.cells[1], s_bg)
        p = row.cells[1].paragraphs[0]
        run = p.add_run(status)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = s_fg

        # Col 2 — note
        p = row.cells[2].paragraphs[0]
        run = p.add_run(note)
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def make_submission_block(doc, num, title, status_str, owner, body_lines):
    """Render a single submission section block."""
    s_bg, s_fg = STATUS.get(status_str, (LIGHT_BG, MUTED))

    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cell = tbl.cell(0, 0)
    set_cell_bg(hdr_cell, LIGHT_BG)
    p = hdr_cell.paragraphs[0]
    r1 = p.add_run(f'#{num}  ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = WHITE
    # Inline colored circle hack: just bold num with colored bg
    r1.font.color.rgb = PRIMARY

    r2 = p.add_run(title + '   ')
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = PRIMARY

    r3 = p.add_run(f'[{status_str}]')
    r3.bold = True
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = s_fg

    r4 = p.add_run(f'   [{owner}]')
    r4.font.size = Pt(8.5)
    r4.font.color.rgb = MUTED

    # Body row
    body_cell = tbl.cell(1, 0)
    set_cell_bg(body_cell, s_bg if s_bg != RGBColor(0xff,0xff,0xff) else LIGHT_BG)
    for line in body_lines:
        p = body_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if line.startswith('- '):
            p.style = doc.styles['List Bullet']
            p.paragraph_format.left_indent = Cm(0.4)
            run = p.add_run(line[2:])
        elif line.startswith('• '):
            p.style = doc.styles['List Bullet']
            p.paragraph_format.left_indent = Cm(0.4)
            run = p.add_run(line[2:])
        else:
            run = p.add_run(line)
        run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════
# COVER / JUDUL
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('RADAR Pangan')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Laporan Progress Tim & Status Submission ke-2')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run('Hackathon Bank Indonesia — Digdaya 2026')
run.font.size = Pt(10)
run.font.color.rgb = MUTED

# Meta table
meta_tbl = doc.add_table(rows=3, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Tanggal Laporan', '14 Mei 2026'),
    ('Deadline Submission ke-2', '4 Juni 2026'),
    ('Sisa Waktu', '21 Hari'),
    ('Branch Aktif', 'feat/ml-training'),
    ('Tim', 'Enzi · Fariz · Rayyan · Rayhan'),
    ('Repository', 'bi-hackathon-group-1'),
]
meta_tbl = doc.add_table(rows=len(meta_data), cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(meta_data):
    row = meta_tbl.rows[i]
    set_cell_bg(row.cells[0], LIGHT_BG)
    set_cell_bg(row.cells[1], RGBColor(0xff,0xff,0xff))
    p = row.cells[0].paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    p = row.cells[1].paragraphs[0]
    r = p.add_run(value)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = PRIMARY

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. RINGKASAN STATUS
# ══════════════════════════════════════════════════════════

heading1(doc, '1. Ringkasan Status Keseluruhan')

# Stat boxes as a 4-col table
stat_tbl = doc.add_table(rows=1, cols=4)
stat_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
stats = [
    ('12', 'Seksi Selesai / Sebagian',   BG_GREEN,  GREEN),
    ('9',  'Seksi Belum Dikerjakan',      BG_RED,    RED_COL),
    ('4',  'Endpoint API Belum Dibuat',   BG_YELLOW, YELLOW),
    ('21', 'Hari Tersisa',                BG_BLUE,   PRIMARY),
]
for i, (num, label, bg, fg) in enumerate(stats):
    c = stat_tbl.rows[0].cells[i]
    set_cell_bg(c, bg)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = c.paragraphs[0]
    r = p.add_run(num + '\n')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = fg
    r2 = p.add_run(label)
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED

doc.add_paragraph().paragraph_format.space_after = Pt(8)

alert_box(doc,
    '⚠ KRITIS: Endpoint /api/v1/simulate WAJIB ada sebelum demo — proposal secara eksplisit menyebutkan "Demo simulasi: skenario sebelum vs. sesudah intervensi". Belum dibangun.',
    BG_RED, 'C0392B')

alert_box(doc,
    '⚡ PERHATIAN: 9 seksi dari 25 seksi submission ke-2 belum dikerjakan sama sekali. Sebagian besar ada di Kelompok F (Pasar & Adopsi) dan di bagian Status Terkini.',
    BG_YELLOW, 'B5830A')

alert_box(doc,
    '✅ KABAR BAIK: ML pipeline 3-lapis sudah jalan, CUSUM ditambahkan, 347.550 baris data historis sudah masuk Supabase, auth sistem (JWT + RBAC) sudah ada, dan FastAPI sudah punya 7 endpoint aktif.',
    BG_GREEN, '2D6A4F')

# ══════════════════════════════════════════════════════════
# 2. PROGRESS PER ANGGOTA TIM
# ══════════════════════════════════════════════════════════

heading1(doc, '2. Progress per Anggota Tim')

# ── ENZI ──────────────────────────────────────────────────
heading2(doc, '2.1  Enzi — ML Core System (Team Lead)')
body_para(doc, 'S2 Data Science, Eötvös Loránd University · Pemilik: semua komponen ML (Lapis 1/2/3, feature engineering, FastAPI)', color=MUTED, size=9)
body_para(doc, 'Progres ML Core: ~70%', size=9)

enzi_data = [
    ('4 model LightGBM (Q50/Q90 × T7/T14)', '✅ Selesai', 'MAPE terbaik 3.52% (val), semua lolos target ≤12%'),
    ('Feature engineering — 28 fitur (lag, rolling, kalender, HET, Z-score)', '✅ Selesai', 'HET coverage 76.2% dari data'),
    ('Lapis 2 — Deteksi: HET threshold + Z-score online', '✅ Selesai', ''),
    ('Lapis 2 — CUSUM (early warning drift detection)', '✅ Selesai', 'Tambahan terbaru (10 Mei); alarm = yellow alert'),
    ('Lapis 2 — Disparity scoring antar kota', '✅ Selesai', ''),
    ('Lapis 3 — LLM ReAct Agent (OpenRouter · Gemini 2.5 Flash)', '✅ Selesai', 'Sebelumnya OpenAI, sudah pindah ke OpenRouter'),
    ('FastAPI inference server — 7 endpoint aktif (port 8001)', '✅ Selesai', '/health, /analyze, /batch, /alerts, /komoditas, /kota, /summary'),
    ('Integrasi Supabase sebagai satu-satunya sumber data ML', '✅ Selesai', '174.290 baris di mart_modelling_harga_pangan'),
    ('report.html — laporan teknis 13 seksi', '✅ Selesai', 'Termasuk Section 13 (ML Roadmap)'),
    ('GET /api/v1/trend/{komoditas}/{kota}', '❌ Belum', 'Riwayat + pita P50/P90 untuk grafik tren'),
    ('GET /api/v1/risk-map/{tanggal}', '❌ Belum', 'Agregasi risiko per provinsi untuk peta koropleth'),
    ('POST /api/v1/simulate', '⚠ WAJIB per proposal', '"skenario sebelum vs. sesudah intervensi"'),
    ('GET/PUT /api/v1/config/het', '❌ Belum', 'Admin config untuk ubah nilai HET'),
    ('Filter & pagination pada GET /api/v1/alerts', '❌ Belum', ''),
    ('Conformal Prediction wrapper', '❌ Belum', 'Proposal menyebut ini tapi belum diimplementasi'),
    ('Fitur is_harvest_season (musim panen per komoditas)', '❌ Belum', 'Hanya Ramadan + akhir tahun yang ada sekarang'),
]
make_status_table(doc, enzi_data)

# ── RAYHAN ────────────────────────────────────────────────
heading2(doc, '2.2  Rayhan — Cloud, Data Ingestion, ETL & Infrastructure')
body_para(doc, 'AWS Certified Cloud Practitioner · Pemilik: Supabase, ETL pipeline, deployment, data validation', color=MUTED, size=9)
body_para(doc, 'Progres Infrastructure: ~60%', size=9)

rayhan_data = [
    ('Supabase PostgreSQL — 4 schema (raw, staging, marts, app)', '✅ Selesai', 'Selesai 6 Mei'),
    ('ETL Pipeline — Postgres loader, dbt, 2 DAG Airflow', '✅ Selesai', 'data_ready_modelling + data_ready_dashboard'),
    ('Load data historis PIHPS — 347.550 baris (2020–2026)', '✅ Selesai', 'Jabar + DKI Jakarta'),
    ('Integrasi cuaca Open-Meteo — 11.605 baris (5 lokasi)', '✅ Selesai', 'Selesai 9 Mei'),
    ('Revisi scope MVP: 6 komoditas bawang+cabai, +Banten +Sulsel', '✅ Selesai', 'constants.py, dbt filter sudah di-update'),
    ('Docker setup + docker-compose', '✅ Selesai', ''),
    ('Auth sistem — JWT + bcrypt + SQLite', '✅ Selesai', 'Selesai 6 Mei — login, admin panel, RBAC'),
    ('dbt re-run dengan filter komoditas (turunkan DB dari ~363 MB)', '❌ Belum', 'DB sekarang 363 MB dari 500 MB limit — harus segera!'),
    ('Load PIHPS historis untuk Banten + Sulawesi Selatan', '❌ Belum', 'Estimasi +125K–375K baris, 2–4 jam proses'),
    ('Validasi data (flag harga hilang, outlier, duplikasi)', '❌ Belum', 'Rencana tapi belum dibangun'),
    ('Deployment cloud (Fly.io / Railway) dengan URL publik', '❌ Belum', 'Juri butuh bisa akses demo'),
    ('HTTPS + CORS untuk API', '❌ Belum', ''),
    ('Diagram arsitektur sistem', '❌ Belum', 'Wajib untuk lampiran submission ke-2'),
]
make_status_table(doc, rayhan_data)

# ── FARIZ + RAYYAN ────────────────────────────────────────
heading2(doc, '2.3  Fariz + Rayyan — RCA Engine & UI')
body_para(doc, 'Bekerja sebagai pasangan: Fariz (domain, produk, proposal) + Rayyan (kuantitatif, implementasi UI, riset data)', color=MUTED, size=9)
body_para(doc, 'Progres RCA + UI: ~30%', size=9)

fr_data = [
    ('File rca_engine.py — ada di src/engine/', '❓ Ada, belum diverifikasi', 'Perlu cek apakah sudah konsumsi output ML API yang sesungguhnya'),
    ('File het_monitor.py — ada di src/engine/', '❓ Ada, belum diverifikasi', 'Perlu cek kesesuaian dengan HET reference terbaru'),
    ('Halaman UI: index.html, login.html, admin.html, debug.html', '🔄 Ada, sebagian', 'Login + admin selesai (Rayhan). index + debug perlu dicek'),
    ('RCA Engine mengonsumsi output ML API (FastAPI port 8001)', '❌ Belum terkonfirmasi', 'Ini inti dari kontribusi Fariz+Rayyan'),
    ('RCA: definisi root cause (supply shock vs seasonal vs HET breach)', '❌ Belum', 'Fariz harus mendefinisikan logika ini dari perspektif domain'),
    ('RCA: narasi rekomendasi yang terbaca oleh TPID/Bapanas', '❌ Belum', '"kemungkinan karena hujan lebat → rekomendasi: operasi pasar"'),
    ('UI terhubung ke endpoint FastAPI (data real, bukan dummy)', '❌ Belum terkonfirmasi', ''),
    ('Penulisan seksi proposal non-teknis (#3, #6, #7, #10, #17, #19, #21, #23)', '❌ Belum', '8 seksi penting, sebagian besar belum disentuh'),
    ('Kompilasi "Evidence of Demand" (BPS CPI, frekuensi pelanggaran HET)', '❌ Belum', 'Rayyan perlu query raw.harga_pangan'),
]
make_status_table(doc, fr_data)

# ══════════════════════════════════════════════════════════
# 3. STATUS 25 SEKSI
# ══════════════════════════════════════════════════════════

doc.add_page_break()
heading1(doc, '3. Status 25 Seksi Submission ke-2')
body_para(doc, 'Submission ke-2 memiliki 25 seksi yang dinilai berdasarkan 6 kriteria juri. Berikut status dan panduan isi setiap seksi.')

# ── KELOMPOK A ────────────────────────────────────────────
heading2(doc, 'Kelompok A — Masalah & Solusi (#1–#4)')
body_para(doc, 'Seksi ini mendefinisikan masalah dan pendekatan solusi. Juri mencari bahasa yang ketat dan spesifik.', color=MUTED, size=9)

make_submission_block(doc, 1, 'Identitas Tim', '✅ Selesai', 'Enzi', [
    'Profil tim, latar belakang masing-masing anggota. Sudah ada di submission pertama.',
    '✅ Tidak perlu perubahan besar — hanya perbarui kalau ada perubahan peran.',
])
make_submission_block(doc, 2, 'Penyempurnaan Masalah & Keselarasan', '🔄 Perlu Update', 'Fariz + Rayyan', [
    'Perketat rumusan masalah. Maks ~180 kata. Harus menjawab: "Apa yang persis rusak di rantai pasok pangan, dan apa yang terjadi kalau tidak diselesaikan?"',
    'Scope sudah direvisi (6 komoditas bawang+cabai, Jabodetabek + Jabar + Sulsel) — framing masalah perlu disesuaikan.',
    '- Sertakan konsekuensi nyata: "TPID bergantung laporan manual 3–5 hari, harga bisa naik 20–30% dalam 48 jam"',
    '- Fokuskan pada komoditas yang dipilih: bawang merah, bawang putih, 4 jenis cabai',
])
make_submission_block(doc, 3, 'Keselarasan Ekosistem', '❌ Belum', 'Fariz', [
    'Siapa pemangku kepentingan dan regulasi apa yang mendukung? Ini adalah "policy hook" — sambungan ke konteks hukum dan kelembagaan.',
    '- Bapanas (otoritas pangan nasional)',
    '- TPID (Tim Pengendalian Inflasi Daerah, 34 provinsi)',
    '- Kemenkop UKM — Koperasi Desa Merah Putih sebagai kanal intervensi (Perpres 9/2025)',
    '- Bulog & ID Food sebagai mitra distribusi',
    '- Relevansi dengan mandat BI dalam pengendalian inflasi pangan volatile',
])
make_submission_block(doc, 4, 'Pendekatan & Mekanisme Solusi', '🔄 Perlu Update', 'Enzi + Fariz', [
    'Maks ~250 kata. Jelaskan sistem 3-lapis + RCA engine. CUSUM dan OpenRouter baru ditambahkan — deskripsi perlu diperbarui.',
    '- Lapis 1: Prediksi harga H+7 dan H+14 dengan LightGBM Quantile (Q50 + Q90)',
    '- Lapis 2: Deteksi — HET threshold, Z-score (spike), CUSUM (drift awal), disparity antar kota',
    '- Lapis 3: LLM ReAct Agent → narasi rekomendasi intervensi dalam Bahasa Indonesia',
    '- RCA Engine: mengubah sinyal ML menjadi analisis akar masalah yang terbaca oleh TPID',
])

# ── KELOMPOK B ────────────────────────────────────────────
heading2(doc, 'Kelompok B — Dampak (#5–#7)')
body_para(doc, 'Seberapa besar dampaknya dan bagaimana cara mengukurnya?', color=MUTED, size=9)

make_submission_block(doc, 5, 'Skala Dampak & Target', '🔄 Perlu Angka', 'Fariz + Rayyan', [
    'Sebutkan skala: 514 kab/kota secara nasional, 280 juta penduduk terdampak. Tapi MVP hanya mencakup 4 provinsi. Jujurlah tentang rollout bertahap.',
    '- MVP saat ini: Jabodetabek + Jawa Barat + Sulawesi Selatan (4 provinsi)',
    '- Target Tahap 2: 10 provinsi prioritas rawan inflasi pangan volatile',
    '- Target jangka panjang: Seluruh 514 kab/kota (via PIHPS yang sudah cover nasional)',
])
make_submission_block(doc, 6, 'Pengukuran Dampak', '❌ Belum', 'Rayyan', [
    'KPI spesifik dan terukur dengan baseline. Ini adalah seksi paling kuantitatif — Rayyan yang paling cocok mengerjakannya.',
    '- MAPE ≤ 12% → sudah tercapai: 3.52% (Q50 T7) ✅',
    '- Kecepatan deteksi pelanggaran HET ≤ 24 jam',
    '- Pengurangan disparitas harga antar kota sebesar 15%',
    '- Perlu sumber: data BPS CPI pangan volatile 2020–2024',
    '- Hitung berapa kali HET dilanggar di raw.harga_pangan → ini "baseline" yang ingin kita perbaiki',
])
make_submission_block(doc, 7, 'Nilai Sistem & Publik', '❌ Belum', 'Fariz', [
    'Jawab: "Nilai sistemik apa yang tercipta melampaui pengguna langsung?"',
    '- Respons pemerintah lebih cepat → volatilitas inflasi pangan lebih rendah → pengurangan kemiskinan',
    '- Hubungkan ke mandat BI dalam pengendalian inflasi',
    '- Keselarasan dengan program Koperasi Desa Merah Putih (Perpres 9/2025) sebagai kanal distribusi',
])

# ── KELOMPOK C ────────────────────────────────────────────
heading2(doc, 'Kelompok C — Inovasi (#8–#10)')
body_para(doc, 'Apa yang membuat solusi ini berbeda dan baru?', color=MUTED, size=9)

make_submission_block(doc, 8, 'Orisinalitas Solusi', '🔄 Sebagian', 'Enzi + Fariz', [
    'Bandingkan secara eksplisit dengan sistem yang ada:',
    '- PIHPS: menampilkan harga, tidak ada prediksi atau rekomendasi',
    '- TaniHub: rantai pasok, bukan kebijakan harga',
    '- EWS BI (makro): level makro, bukan per komoditas per kota',
    'Diferensiator RADAR Pangan: prediksi real-time + RCA engine + rekomendasi intervensi dalam satu sistem.',
])
make_submission_block(doc, 9, 'Inovasi Teknologi & Metode', '🔄 Perlu Update', 'Enzi', [
    '- LightGBM Quantile Regression — prediksi dengan interval ketidakpastian (bukan hanya titik)',
    '- CUSUM — mendeteksi drift harga sebelum menjadi krisis (early warning 2–4 hari lebih awal)',
    '- Z-score online — mendeteksi lonjakan tiba-tiba',
    '- LLM ReAct Agent (Gemini 2.5 Flash via OpenRouter) — narasi intervensi dalam Bahasa Indonesia',
    '- Conformal Prediction — CATATAN: direncanakan Tahap 2, belum diimplementasi; harus disebutkan jujur',
])
make_submission_block(doc, 10, 'Kreativitas dalam Implementasi', '❌ Belum', 'Fariz + Rayyan', [
    'Ini tentang BAGAIMANA Anda membangunnya, bukan APA yang dibangun. Dua elemen genuinely kreatif:',
    '- RCA Engine: mengubah sinyal ML menjadi narasi akar masalah terbaca TPID — "kemungkinan besar disebabkan hujan lebat di Cirebon + musim panen berakhir → rekomendasi: operasi pasar"',
    '- Koperasi Desa Merah Putih sebagai kanal intervensi — memanfaatkan Perpres 9/2025, jaringan 70.000+ koperasi sebagai "last-mile" distribusi',
    '- UI berbasis peran: tampilan berbeda untuk petugas TPID, admin Bapanas, dan viewer publik',
])

# ── KELOMPOK D ────────────────────────────────────────────
doc.add_page_break()
heading2(doc, 'Kelompok D — Kelayakan Teknis (#11–#14)')
body_para(doc, 'Apakah ini bisa benar-benar bekerja di skala besar?', color=MUTED, size=9)

make_submission_block(doc, 11, 'Arsitektur Sistem', '🔄 Sebagian', 'Rayhan + Enzi', [
    'Rayhan WAJIB membuat diagram arsitektur. Alur yang harus ditampilkan:',
    'PIHPS API → ETL (Postgres Loader + dbt) → Supabase PostgreSQL → ML API (FastAPI :8001) → RCA Engine → Frontend (Alpine.js) + Auth Layer',
    'report.html sudah ada ✅, tapi diagram visual diperlukan sebagai lampiran. Gunakan draw.io / Excalidraw / Figma.',
])
make_submission_block(doc, 12, 'Data & Kelayakan', '🔄 Sebagian', 'Rayhan + Enzi', [
    'Data tersedia: PIHPS ✅, HET reference ✅, cuaca Open-Meteo ✅.',
    '- BPS Statistik Tanaman Sayuran (produksi per provinsi) → proxy untuk fitur stok relatif — BELUM ADA',
    '- BI SEKI Tabel VII — CPI pangan volatile bulanan → fitur makro — BELUM ADA',
    '- Data stok Koperasi Desa Merah Putih → tidak tersedia publik (nyatakan jujur di proposal)',
])
make_submission_block(doc, 13, 'Keamanan & Kepatuhan', '🔄 Sebagian', 'Rayhan', [
    'JWT + bcrypt sudah dibangun ✅. Yang masih perlu ditambahkan:',
    '- Catatan UU PDP/PDPA: tidak ada data pribadi yang disimpan, hanya data agregat harga',
    '- Data residency: Supabase di-host di region aws-1-ap-northeast-1 (Tokyo)',
    '- Dokumentasi RBAC: admin / analyst / viewer — hak akses masing-masing',
    '- PERHATIAN: SECRET_KEY JWT masih hardcoded — harus dipindahkan ke env variable sebelum demo',
])
make_submission_block(doc, 14, 'Kesiapan Implementasi (MVP)', '🔄 Sebagian', 'Semua', [
    '5 halaman ada di frontend. ML API punya 7 endpoint aktif.',
    '- Nyatakan dengan jelas apa yang sudah siap demo vs yang direncanakan',
    '- Endpoint /simulate belum dibangun — juri akan memeriksa ini',
    '- Apakah UI sudah mengonsumsi data real dari FastAPI? Perlu dikonfirmasi',
])

# ── KELOMPOK E ────────────────────────────────────────────
heading2(doc, 'Kelompok E — Model Bisnis (#15–#18)')
body_para(doc, 'Bagaimana sistem ini bisa berkelanjutan secara finansial?', color=MUTED, size=9)

make_submission_block(doc, 15, 'Proposisi Nilai', '🔄 Perlu Update', 'Fariz + Rayyan', [
    'Revisi untuk scope baru (6 komoditas, Jabodetabek + Jabar + Sulsel). Format: Masalah → Solusi → Manfaat.',
    '- Untuk TPID: "Bisa merespons pelanggaran HET 24 jam lebih cepat"',
    '- Untuk Bapanas: "Bisa memprioritaskan wilayah mana yang butuh operasi pasar"',
    '- Maks ~220 kata',
])
make_submission_block(doc, 16, 'Pendapatan & Pendanaan', '🔄 Sebagian', 'Fariz', [
    '- B2G SaaS: langganan pemerintah — Bapanas, 34 TPID Provinsi. Estimasi: Rp 50–200 juta/tahun per provinsi',
    '- B2B Analytics: Bulog, ID Food, retailer pangan besar',
    '- Perlu menambahkan model pricing yang spesifik',
])
make_submission_block(doc, 17, 'Struktur Biaya & Keberlanjutan', '❌ Belum', 'Fariz + Rayhan', [
    'Estimasi biaya operasional yang diperlukan juri:',
    '- Cloud (Supabase): Free tier sekarang → ~$25–50/bulan tier berbayar',
    '- Deployment (Fly.io/Railway): ~$10–30/bulan',
    '- LLM API (OpenRouter · Gemini 2.5 Flash): ~$0.10–0.50 per analisis tergantung volume',
    '- Tim pengembang: estimasi biaya maintenance',
    '- Titik impas: berapa pelanggan B2G yang dibutuhkan untuk menutup biaya?',
])
make_submission_block(doc, 18, 'Skalabilitas', '🔄 Sebagian', 'Rayhan + Fariz', [
    '- Teknis: Docker + cloud = scale horizontal ✅',
    '- Data: Supabase free tier 500 MB, perlu rencana migrasi BigQuery/GCS',
    '- Produk: 70.000+ Koperasi Desa Merah Putih sebagai titik distribusi (Perpres 9/2025)',
    '- Perlu angka rollout: "Tahun 1: 4 provinsi → Tahun 2: 10 provinsi → Tahun 3: 34 provinsi"',
])

# ── KELOMPOK F ────────────────────────────────────────────
doc.add_page_break()
heading2(doc, 'Kelompok F — Pasar & Kesiapan Adopsi (#19–#23)')
body_para(doc, 'Siapa yang akan memakai ini dan apakah mereka benar-benar akan mengadopsinya?', color=MUTED, size=9)

make_submission_block(doc, 19, 'Kemitraan & Distribusi', '❌ Belum', 'Fariz', [
    'Sebutkan nama mitra potensial (meski belum ada MOU, bisa sebut "dalam diskusi" atau "target mitra"):',
    '- Bapanas — otoritas pangan nasional',
    '- TPID 34 Provinsi — pengguna utama',
    '- Kemenkop UKM — kanal koperasi',
    '- Bulog — logistik dan distribusi',
    '- ID Food — holding BUMN pangan',
    '- Kanal distribusi: pengadaan pemerintah via e-Katalog LKPP',
])
make_submission_block(doc, 20, 'Kesesuaian Masalah-Pasar', '🔄 Perlu Diperketat', 'Fariz + Rayyan', [
    'Maks ~120 kata. Konsekuensi spesifik kalau tidak diselesaikan:',
    '"Tanpa sistem ini, TPID bergantung pada laporan manual yang membutuhkan 3–5 hari untuk terkumpul, sementara harga bisa naik 20–30% dalam 48 jam saat ada shock supply. Akibatnya, operasi pasar dilakukan terlambat dan tidak tepat sasaran."',
])
make_submission_block(doc, 21, 'Bukti Permintaan (Evidence of Demand)', '❌ Belum', 'Rayyan', [
    'Harus menunjukkan bahwa masalah ini nyata dengan data. Rayyan perlu mengerjakan ini:',
    '- BPS CPI pangan volatile: bawang merah/cabai sering mengalami inflasi tahunan 15–40%',
    '- PIHPS sudah cover 514 kab/kota → pemerintah sudah investasi di pemantauan harga',
    '- Frekuensi pelanggaran HET: query raw.harga_pangan dan bandingkan dengan het_reference',
    '- BI SEKI Tabel VII — data CPI pangan volatile bulanan sebagai referensi tambahan',
])
make_submission_block(doc, 22, 'Target Pasar', '🔄 Sebagian', 'Fariz', [
    'Segmentasi pasar dengan ukuran masing-masing:',
    '- Nasional: Bapanas (1 entitas, nilai tinggi, impact luas)',
    '- Provinsi: 34 TPID unit — pasar utama B2G',
    '- Institusional: Bulog, ID Food, korporasi pangan besar',
    '- Potensi pendapatan: 34 × Rp 50 juta/tahun = Rp 1,7 miliar ARR',
])
make_submission_block(doc, 23, 'Kesiapan Adopsi', '❌ Belum', 'Fariz + Rayyan', [
    'Hambatan adopsi dan cara mengatasinya — juri mencari realisme:',
    '- Hambatan 1: Akses internet di daerah terpencil → solusi: mode offline untuk dashboard',
    '- Hambatan 2: Pelatihan operator pemerintah → solusi: program onboarding, UI sederhana berbasis peran',
    '- Hambatan 3: Kepercayaan pada rekomendasi AI → solusi: posisikan AI sebagai "co-pilot, bukan pengambil keputusan"',
    '- Hambatan 4: Proses pengadaan pemerintah yang lambat → solusi: daftarkan di e-Katalog LKPP',
])

# ── KELOMPOK G ────────────────────────────────────────────
heading2(doc, 'Kelompok G — Status Terkini & Lampiran (#24–#25 + Lampiran)')
body_para(doc, 'Di mana Anda sekarang dan apa yang bisa dilihat juri?', color=MUTED, size=9)

make_submission_block(doc, 24, 'Progress Sejak Submission ke-1', '❌ Belum', 'Enzi', [
    'Ini adalah "quick win" terbesar — Anda sudah punya banyak yang bisa dilaporkan. Buat daftar konkret:',
    '- ✅ 4 model LightGBM dilatih, MAPE terbaik 3.52%',
    '- ✅ Database cloud Supabase live (347.550+ baris, data cuaca terintegrasi)',
    '- ✅ Pipeline ML 3-lapis berjalan (Lapis 1/2/3)',
    '- ✅ CUSUM early-warning detection ditambahkan',
    '- ✅ OpenRouter LLM agent (Gemini 2.5 Flash) terintegrasi',
    '- ✅ Sistem auth (JWT + RBAC: admin/analyst/viewer)',
    '- ✅ FastAPI 7 endpoint aktif',
    '- ✅ Laporan teknis (report.html) — 13 seksi',
    '- ✅ Integrasi cuaca Open-Meteo (5 lokasi, 11.605 baris)',
    '- 🔄 RCA Engine — dalam pengerjaan (Fariz + Rayyan)',
])
make_submission_block(doc, 25, 'Status Terkini Sistem', '❌ Belum', 'Enzi', [
    'Tahap saat ini: PROTOTIPE (API ML berjalan, frontend sebagian, database live). Juri menghargai kejujuran.',
    '- "ML inference API berjalan lokal dan cloud deployment sedang dalam proses"',
    '- "Halaman frontend ada tapi sebagian belum mengonsumsi semua endpoint API"',
    '- "Endpoint /simulate sedang dibangun — target selesai [tanggal]"',
    '- Sertakan URL demo publik (begitu Rayhan selesai deploy)',
    '- Sertakan link ke report.html',
])
make_submission_block(doc, 'LAM', 'Lampiran', '🔄 Sebagian', 'Semua', [
    '- ✅ report.html — laporan teknis lengkap (13 seksi)',
    '- ❌ Diagram arsitektur — Rayhan perlu membuat (draw.io / Excalidraw / Figma)',
    '- ❌ Video demo atau screenshot — tampilkan CUSUM alert + rekomendasi LLM dalam aksi',
    '- ❌ URL demo publik — Rayhan perlu deploy ke cloud',
])

# ══════════════════════════════════════════════════════════
# 4. RENCANA KERJA 21 HARI
# ══════════════════════════════════════════════════════════

doc.add_page_break()
heading1(doc, '4. Rencana Kerja 21 Hari (14 Mei – 3 Juni 2026)')

alert_box(doc,
    'Deadline submission ke-2: 4 Juni 2026. Target selesai semua pekerjaan pada 3 Juni 2026 (buffer 1 hari untuk review akhir).',
    BG_BLUE, '1a3a5c')

heading3(doc, 'Minggu 1: 14–21 Mei — Celah Kritis')
for item in [
    ('[Enzi] Bangun endpoint POST /api/v1/simulate — WAJIB per proposal',),
    ('[Rayhan] Re-run dbt dengan filter komoditas → turunkan DB dari 363 MB ke ~175 MB',),
    ('[Rayhan] Load PIHPS historis untuk Banten + Sulawesi Selatan',),
    ('[Fariz+Rayyan] Tulis seksi #3 (Keselarasan Ekosistem), #7 (Nilai Publik), #10 (Kreativitas)',),
    ('[Fariz+Rayyan] Konfirmasi status RCA engine: apakah sudah konsumsi output FastAPI?',),
]:
    bullet(doc, item[0])

heading3(doc, 'Minggu 2: 22–28 Mei — Konten Inti')
for item in [
    '[Rayhan] Buat diagram arsitektur sistem (draw.io/Excalidraw)',
    '[Rayhan] Deploy FastAPI ke cloud (Fly.io/Railway), buat URL publik demo',
    '[Enzi] Tulis seksi #24 (Progress Sejak Submission ke-1) dan #25 (Status Terkini)',
    '[Enzi] Update seksi #9 (Inovasi Teknologi) — sertakan CUSUM, OpenRouter',
    '[Enzi] Bangun endpoint GET /api/v1/trend/{komoditas}/{kota}',
    '[Fariz+Rayyan] Tulis seksi #6 (Pengukuran Dampak), #17 (Struktur Biaya), #19 (Kemitraan)',
    '[Rayyan] Kompilasi "Evidence of Demand" (#21) — query Supabase untuk hitung pelanggaran HET',
]:
    bullet(doc, item)

heading3(doc, 'Minggu 3: 29 Mei – 3 Juni — Poles & Finalisasi')
for item in [
    '[Semua] Review semua 25 seksi, pastikan ada di batas kata yang ditentukan',
    '[Fariz+Rayyan] Finalisasi halaman UI — pastikan semua data dari API real',
    '[Semua] Buat video demo singkat (2–3 menit): CUSUM alert → LLM rekomendasi → RCA narasi',
    '[Rayhan] Pastikan URL demo publik stabil, HTTPS aktif',
    '[Enzi] Final review report.html — update kalau ada perubahan teknis',
    '[Semua] Kumpulkan semua lampiran: diagram, demo video, report.html',
]:
    bullet(doc, item)

# Matriks prioritas per orang
heading2(doc, 'Matriks Prioritas per Anggota Tim')

prio_tbl = doc.add_table(rows=1, cols=4)
prio_tbl.style = 'Table Grid'
prio_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers  = ['Enzi (ML)', 'Rayhan (Infrastructure)', 'Fariz (Domain+Proposal)', 'Rayyan (Kuantitatif+UI)']
colors   = [PRIMARY, GREEN, ACCENT, ACCENT]
tasks    = [
    [
        '1. POST /api/v1/simulate — WAJIB',
        '2. Tulis seksi #24 dan #25',
        '3. GET /api/v1/trend/{komoditas}/{kota}',
        '4. Update seksi #9 (inovasi teknologi)',
        '5. GET /api/v1/risk-map/{tanggal}',
        '6. Conformal Prediction (atau nyatakan "Tahap 2")',
        '7. Fitur is_harvest_season di features.py',
    ],
    [
        '1. dbt re-run dengan filter komoditas',
        '2. Load PIHPS Banten + Sulsel',
        '3. Buat diagram arsitektur',
        '4. Deploy FastAPI ke cloud (URL publik)',
        '5. Tambah validasi data',
        '6. Set HTTPS + CORS untuk API',
        '7. Tulis seksi #13 (Keamanan & Kepatuhan)',
    ],
    [
        '1. Tulis seksi #3 (Keselarasan Ekosistem)',
        '2. Tulis seksi #7 (Nilai Sistem & Publik)',
        '3. Tulis seksi #10 (Kreativitas Implementasi)',
        '4. Tulis seksi #19 (Kemitraan & Distribusi)',
        '5. Tulis seksi #17 (Struktur Biaya)',
        '6. Definisikan logika RCA: "operasi pasar" vs "redistribusi stok"',
        '7. Tulis seksi #23 (Kesiapan Adopsi)',
    ],
    [
        '1. Kompilasi "Evidence of Demand" (#21)',
        '2. Query Supabase: hitung pelanggaran HET per komoditas',
        '3. Tulis seksi #6 (Pengukuran Dampak + KPI numerik)',
        '4. Riset BPS CPI pangan volatile 2020–2024',
        '5. Implementasi UI konsumsi FastAPI',
        '6. Implementasi RCA engine bersama Fariz',
        '7. Cek BI SEKI Tabel VII untuk data CPI volatile',
    ],
]

hdr_cells = prio_tbl.rows[0].cells
for i, (cell, hdr, col) in enumerate(zip(hdr_cells, headers, colors)):
    set_cell_bg(cell, col)
    p = cell.paragraphs[0]
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = WHITE

body_row = prio_tbl.add_row()
for i, (cell, task_list) in enumerate(zip(body_row.cells, tasks)):
    set_cell_bg(cell, LIGHT_BG)
    first = True
    for t in task_list:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(t)
        r.font.size = Pt(8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ── Tabel 10 Seksi Belum Disentuh ─────────────────────────
heading2(doc, 'Ringkasan: 10 Seksi yang Belum Disentuh')

missing_data = [
    ('#3',  'Keselarasan Ekosistem',            'Fariz',           'Policy hook — tanpa ini proposal terlihat tidak terhubung ke regulasi'),
    ('#6',  'Pengukuran Dampak',                'Rayyan',          'KPI numerik — juri butuh angka konkret untuk menilai dampak'),
    ('#7',  'Nilai Sistem & Publik',            'Fariz',           'Koneksi ke dampak sistemik yang lebih besar'),
    ('#10', 'Kreativitas dalam Implementasi',   'Fariz + Rayyan',  'RCA engine dan kanal Koperasi adalah elemen paling kreatif'),
    ('#17', 'Struktur Biaya & Keberlanjutan',   'Fariz + Rayhan',  'Juri ingin tahu apakah bisnis ini viable secara finansial'),
    ('#19', 'Kemitraan & Distribusi',           'Fariz',           'Tanpa mitra distribusi, adopsi tidak bisa terjadi'),
    ('#21', 'Bukti Permintaan',                 'Rayyan',          'Data BPS + frekuensi pelanggaran HET — membuktikan masalah nyata ada'),
    ('#23', 'Kesiapan Adopsi',                  'Fariz + Rayyan',  'Juri mencari realisme — apa hambatan dan bagaimana mengatasinya'),
    ('#24', 'Progress Sejak Submission ke-1',   'Enzi',            '"Quick win" terbesar — kita punya banyak yang bisa dilaporkan'),
    ('#25', 'Status Terkini Sistem',            'Enzi',            'Juri ingin tahu: di mana sistem sekarang dan apa yang bisa diakses?'),
]

miss_tbl = doc.add_table(rows=1, cols=4)
miss_tbl.style = 'Table Grid'
for cell, label in zip(miss_tbl.rows[0].cells, ['#', 'Nama Seksi', 'Pemilik', 'Mengapa Penting']):
    set_cell_bg(cell, PRIMARY)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = WHITE

for idx, (num, name, owner, why) in enumerate(missing_data):
    row = miss_tbl.add_row()
    bg = LIGHT_BG if idx % 2 == 1 else RGBColor(0xff,0xff,0xff)
    data = [num, name, owner, why]
    for cell, val in zip(row.cells, data):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)

# ══════════════════════════════════════════════════════════
# FOOTER PARAGRAPH
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('RADAR Pangan — Progress Report Internal · Dibuat 14 Mei 2026 · Tim Hackathon BI Digdaya 2026 (Kelompok 1)')
r.font.size = Pt(8)
r.font.color.rgb = MUTED

# ══════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════
out = r'D:\Enzi-Folder\personal-project\hackathon-project\bi-hackathon-group-1\progress-report.docx'
doc.save(out)
print(f'Saved: {out}')
